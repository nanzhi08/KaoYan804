import os
from pathlib import Path
from uuid import uuid4

import aiofiles
from fastapi import APIRouter, Depends, UploadFile, File
from sqlalchemy import select as sa_select
from sqlalchemy.ext.asyncio import AsyncSession

from ..config import settings
from ..database import get_db
from ..schemas.common import APIResponse

router = APIRouter(prefix="/api/documents", tags=["文档管理"])

_ALLOWED_TYPES = {"pdf", "doc", "docx", "txt", "md", "png", "jpg", "jpeg"}
_MOJIBAKE_CHARS = set("ÄÅÆÇÉÊËÌÍÎÏÐÑÒÓÔÕÖ×ØÙÚÛÜÝÞßàáâãäåæçèéêëìíîïðñòóôõö÷øùúûüýþÿ¶·¸¹º»¼½¾¿")


def _cjk_count(text: str) -> int:
    return sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")


def _mojibake_score(text: str) -> int:
    return sum(1 for ch in text if ch in _MOJIBAKE_CHARS)


def _repair_filename(name: str) -> str:
    if not name:
        return name

    best = name
    best_cjk = _cjk_count(name)
    best_noise = _mojibake_score(name)

    for source_encoding in ("latin1", "cp1252"):
        for target_encoding in ("utf-8", "gb18030", "gbk"):
            try:
                candidate = name.encode(source_encoding).decode(target_encoding)
            except (UnicodeEncodeError, UnicodeDecodeError):
                continue

            if "\ufffd" in candidate:
                continue

            candidate_cjk = _cjk_count(candidate)
            candidate_noise = _mojibake_score(candidate)

            if candidate_cjk > best_cjk or (
                candidate_cjk == best_cjk and candidate_noise < best_noise
            ):
                best = candidate
                best_cjk = candidate_cjk
                best_noise = candidate_noise

    return best


def _safe_original_name(raw_filename: str) -> str:
    repaired = _repair_filename(raw_filename or "unnamed")
    return Path(repaired).name or "unnamed"


def _stored_filename(original_name: str) -> str:
    suffix = Path(original_name).suffix.lower()
    stem = Path(original_name).stem[:80] or "document"
    safe_stem = "".join(ch if ch.isalnum() or ch in "._-一-龥" else "_" for ch in stem).strip("._")
    return f"{safe_stem or 'document'}_{uuid4().hex[:12]}{suffix}"


def _file_type(filename: str) -> str:
    suffix = Path(filename).suffix.lower().lstrip(".")
    return suffix or "txt"


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _extract_content(file_type: str, content: bytes) -> str:
    if file_type in {"txt", "md"}:
        return _decode_text(content)

    if file_type == "doc":
        return "[.doc 格式不支持文本提取，请先转换为 .docx 后上传。]"

    if file_type == "docx":
        try:
            from io import BytesIO
            from docx import Document

            doc = Document(BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            return f"[解析失败: {e}]"

    if file_type == "pdf":
        try:
            from ..services.ocr_service import ocr_pdf

            _was_scanned, text = ocr_pdf(content)
            return text
        except Exception as e:
            return f"[解析失败: {e}]"

    if file_type in {"png", "jpg", "jpeg"}:
        try:
            from ..services.ocr_service import ocr_image

            return ocr_image(content)
        except Exception as e:
            return f"[解析失败: {e}]"

    return ""


async def _repair_document_record(doc) -> bool:
    fixed_filename = Path(_repair_filename(doc.filename or "")).name
    fixed_original_name = _safe_original_name(doc.original_name or "unnamed")
    changed = False

    if fixed_filename and fixed_filename != doc.filename:
        old_path = os.path.join(settings.UPLOAD_DIR, doc.filename)
        new_path = os.path.join(settings.UPLOAD_DIR, fixed_filename)
        if os.path.exists(old_path) and old_path != new_path and not os.path.exists(new_path):
            os.rename(old_path, new_path)
        doc.filename = fixed_filename
        changed = True

    if fixed_original_name != doc.original_name:
        doc.original_name = fixed_original_name
        changed = True

    return changed


async def _repair_documents_if_needed(db: AsyncSession, docs: list) -> None:
    changed = False
    for doc in docs:
        changed = await _repair_document_record(doc) or changed

    if changed:
        await db.commit()


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    from ..models.document import Document

    content = await file.read()
    original_name = _safe_original_name(file.filename or "unnamed")
    file_type = _file_type(original_name)

    if file_type not in _ALLOWED_TYPES:
        return APIResponse(code=400, message="不支持的文件类型")

    existing = await db.execute(
        sa_select(Document).where(
            Document.original_name == original_name,
            Document.file_size == len(content),
        )
    )
    if existing.scalar_one_or_none():
        return APIResponse(code=409, message="该文件已存在，请勿重复上传")

    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    stored_name = _stored_filename(original_name)
    file_path = os.path.join(settings.UPLOAD_DIR, stored_name)

    async with aiofiles.open(file_path, "wb") as f:
        await f.write(content)

    doc = Document(
        filename=stored_name,
        original_name=original_name,
        file_type=file_type,
        file_size=len(content),
        content_text=_extract_content(file_type, content),
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    return APIResponse(data={"id": doc.id, "filename": original_name, "file_type": file_type})


@router.get("")
async def list_documents(db: AsyncSession = Depends(get_db)):
    from ..models.document import Document

    result = await db.execute(sa_select(Document).order_by(Document.uploaded_at.desc()))
    docs = result.scalars().all()
    await _repair_documents_if_needed(db, docs)
    return APIResponse(data=[
        {
            "id": d.id,
            "original_name": d.original_name,
            "file_type": d.file_type,
            "file_size": d.file_size,
            "tags": d.tags,
            "uploaded_at": d.uploaded_at.isoformat(),
        }
        for d in docs
    ])


@router.get("/{doc_id}")
async def get_document(doc_id: int, db: AsyncSession = Depends(get_db)):
    from ..models.document import Document

    result = await db.execute(sa_select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return APIResponse(code=404, message="文档不存在")
    await _repair_documents_if_needed(db, [doc])
    return APIResponse(data={
        "id": doc.id,
        "original_name": doc.original_name,
        "file_type": doc.file_type,
        "file_size": doc.file_size,
        "tags": doc.tags,
        "content_text": doc.content_text,
        "uploaded_at": doc.uploaded_at.isoformat(),
    })


@router.get("/{doc_id}/content")
async def get_document_content(doc_id: int, db: AsyncSession = Depends(get_db)):
    from ..models.document import Document

    result = await db.execute(sa_select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return APIResponse(code=404, message="文档不存在")
    await _repair_documents_if_needed(db, [doc])
    return APIResponse(data={"content": doc.content_text})


@router.delete("/{doc_id}")
async def delete_document(doc_id: int, db: AsyncSession = Depends(get_db)):
    from ..models.document import Document

    result = await db.execute(sa_select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return APIResponse(code=404, message="文档不存在")
    await _repair_documents_if_needed(db, [doc])

    file_path = os.path.join(settings.UPLOAD_DIR, doc.filename)
    await db.delete(doc)
    await db.commit()

    if os.path.isfile(file_path):
        os.remove(file_path)

    return APIResponse(message="删除成功")
