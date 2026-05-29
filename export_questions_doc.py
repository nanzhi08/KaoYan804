"""导出全部题库为 Word 文档"""
import asyncio
import sys
sys.path.insert(0, "backend")

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from sqlalchemy import select
from app.database import async_session, init_db
from app.models.question import Question, QuestionKnowledgePoint
from app.models.knowledge_point import KnowledgePoint

TYPE_NAMES = {
    "single_choice": "单选题",
    "multi_choice": "多选题",
    "fill_blank": "填空题",
    "program_reading": "程序阅读题",
    "analysis": "分析题",
    "calculation": "计算题",
    "programming": "编程题",
    "short_answer": "简答题",
}

PART_NAMES = {
    "C_programming": "第一部分：C语言高级程序设计",
    "data_structure": "第二部分：数据结构",
}

DIFFICULTY_STARS = {1: "★☆☆☆☆", 2: "★★☆☆☆", 3: "★★★☆☆", 4: "★★★★☆", 5: "★★★★★"}


def set_cell_font(cell, name="微软雅黑", size=10, bold=False):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
            run.font.size = Pt(size)
            run.font.bold = bold


def add_question(doc, q, kp_name):
    """Add a single question to the document"""
    # Question header: type + difficulty
    h = doc.add_heading(level=3)
    run = h.add_run(f"[{q.id}] {TYPE_NAMES.get(q.type, q.type)}  {DIFFICULTY_STARS.get(q.difficulty, '')}")
    run.font.size = Pt(12)

    # Knowledge point + part
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"章节：{kp_name}  |  所属：{PART_NAMES.get(q.part, q.part)}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)

    # Content
    content_para = doc.add_paragraph()
    content_run = content_para.add_run("【题目】")
    content_run.font.bold = True
    content_run.font.size = Pt(10.5)
    content_text = content_para.add_run(f"\n{q.content}")
    content_text.font.size = Pt(10.5)

    # Options (for choice questions)
    if q.options and isinstance(q.options, dict) and len(q.options) > 0:
        opt_para = doc.add_paragraph()
        opt_run = opt_para.add_run("选项：")
        opt_run.font.bold = True
        opt_run.font.size = Pt(10.5)
        for k, v in q.options.items():
            opt_text = opt_para.add_run(f"\n  {k}. {v}")
            opt_text.font.size = Pt(10.5)

    # Code snippet (if any)
    # The code is already embedded in content via markdown code blocks, so we skip duplicate

    # Answer
    ans_para = doc.add_paragraph()
    ans_run = ans_para.add_run("【答案】")
    ans_run.font.bold = True
    ans_run.font.size = Pt(10.5)
    ans_run.font.color.rgb = RGBColor(0, 100, 0)
    ans_text = ans_para.add_run(f"\n{q.answer}")
    ans_text.font.size = Pt(10.5)
    ans_text.font.color.rgb = RGBColor(0, 100, 0)

    # Explanation
    if q.explanation:
        exp_para = doc.add_paragraph()
        exp_run = exp_para.add_run("【解析】")
        exp_run.font.bold = True
        exp_run.font.size = Pt(10.5)
        exp_run.font.color.rgb = RGBColor(0, 70, 140)
        exp_text = exp_para.add_run(f"\n{q.explanation}")
        exp_text.font.size = Pt(10.5)
        exp_text.font.color.rgb = RGBColor(0, 70, 140)

    # Source
    if q.source:
        src_para = doc.add_paragraph()
        src_run = src_para.add_run(f"来源：{q.source}")
        src_run.font.size = Pt(8)
        src_run.font.color.rgb = RGBColor(150, 150, 150)

    # Separator
    doc.add_paragraph("─" * 60)


async def main():
    await init_db()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Title
    title = doc.add_heading('上海第二工业大学 804《数据结构与高级程序设计》题库', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('共计 119 题  |  高级程序设计(70分) + 数据结构(80分)  |  2026年考试大纲')
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    async with async_session() as session:
        # Query all questions with their knowledge points
        result = await session.execute(
            select(Question).order_by(Question.part, Question.id)
        )
        questions = result.scalars().all()

        # Query knowledge points
        kp_result = await session.execute(select(KnowledgePoint))
        kps = {kp.id: kp for kp in kp_result.scalars().all()}

        current_part = None
        count = 0
        stats = {}

        for q in questions:
            # Section header when part changes
            if q.part != current_part:
                current_part = q.part
                section = doc.add_heading(PART_NAMES.get(q.part, q.part), level=1)

            # Get knowledge point name
            kp_name = "未分类"
            if q.knowledge_points:
                for qkp in q.knowledge_points:
                    kp = kps.get(qkp.knowledge_point_id)
                    if kp:
                        kp_name = kp.name
                        break

            add_question(doc, q, kp_name)
            count += 1

            # Stats
            t = q.type
            stats[t] = stats.get(t, 0) + 1

    # Statistics page at the end
    doc.add_page_break()
    doc.add_heading("题库统计", level=1)

    # Summary table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "题型"
    hdr_cells[1].text = "数量"
    hdr_cells[2].text = "所属部分"
    hdr_cells[3].text = "考试中出现"

    for t, cnt in sorted(stats.items()):
        row = table.add_row().cells
        row[0].text = TYPE_NAMES.get(t, t)
        row[1].text = str(cnt)

    part_c = sum(1 for q in questions if q.part == "C_programming")
    part_ds = sum(1 for q in questions if q.part == "data_structure")

    doc.add_paragraph()
    doc.add_paragraph(f"C语言高级程序设计：{part_c} 题")
    doc.add_paragraph(f"数据结构：{part_ds} 题")
    doc.add_paragraph(f"总计：{count} 题")

    output_path = "C:/Users/zhangsihai/Desktop/考研知识库系统/804题库_119题.docx"
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")
    print(f"共导出 {count} 道题目")


if __name__ == "__main__":
    asyncio.run(main())
